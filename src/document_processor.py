"""
Document Processor Module
Handles extraction of text from various document formats (PDF, DOCX, TXT)
"""

import os
import re
from typing import Optional, Dict, Any
from pathlib import Path

try:
    import PyPDF2
    from docx import Document
    import docx2txt
except ImportError as e:
    print(f"Warning: Some document processing libraries not available: {e}")


class DocumentProcessor:
    """Handles text extraction from various document formats."""
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.doc', '.txt'}
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from a document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
        except Exception as e:
            raise Exception(f"Error processing {file_path}: {str(e)}")
    
    def _extract_from_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF file."""
        text = ""
        metadata = {}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'producer': pdf_reader.metadata.get('/Producer', ''),
                        'pages': len(pdf_reader.pages)
                    }
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page_text
                        
        except Exception as e:
            raise Exception(f"PDF extraction failed: {str(e)}")
        
        return {
            'text': text.strip(),
            'metadata': metadata,
            'file_type': 'pdf',
            'file_size': file_path.stat().st_size
        }
    
    def _extract_from_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        try:
            # Try docx2txt first (better for complex documents)
            text = docx2txt.process(str(file_path))
            
            # If docx2txt fails, try python-docx
            if not text.strip():
                doc = Document(file_path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extract metadata
            doc = Document(file_path)
            core_props = doc.core_properties
            metadata = {
                'title': getattr(core_props, 'title', '') or '',
                'author': getattr(core_props, 'author', '') or '',
                'subject': getattr(core_props, 'subject', '') or '',
                # 'creator' is not a standard property in python-docx, fallback to author
                'creator': getattr(core_props, 'author', '') or '',
                'pages': len(doc.sections) if doc.sections else 0
            }
            
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")
        
        return {
            'text': text.strip(),
            'metadata': metadata,
            'file_type': 'docx',
            'file_size': file_path.stat().st_size
        }
    
    def _extract_from_txt(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from TXT file."""
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except UnicodeDecodeError:
            # Fallback to other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise Exception("Could not decode text file with any supported encoding")
        
        return {
            'text': text.strip(),
            'metadata': {
                'title': file_path.stem,
                'author': '',
                'subject': '',
                'creator': '',
                'pages': 1
            },
            'file_type': 'txt',
            'file_size': file_path.stat().st_size
        }
    
    def clean_text(self, text: str) -> str:
        """
        Clean and preprocess extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text ready for summarization
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove page markers
        text = re.sub(r'--- Page \d+ ---', '', text)
        
        # Remove headers/footers (common patterns)
        text = re.sub(r'^\d+\s*$', '', text, flags=re.MULTILINE)
        
        # Remove excessive newlines
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def get_document_stats(self, text: str) -> Dict[str, Any]:
        """
        Get statistics about the document.
        
        Args:
            text: Document text
            
        Returns:
            Dictionary with document statistics
        """
        words = text.split()
        sentences = re.split(r'[.!?]+', text)
        paragraphs = [p for p in text.split('\n\n') if p.strip()]
        
        return {
            'word_count': len(words),
            'sentence_count': len([s for s in sentences if s.strip()]),
            'paragraph_count': len(paragraphs),
            'character_count': len(text),
            'estimated_reading_time_minutes': len(words) / 200  # Average reading speed
        } 